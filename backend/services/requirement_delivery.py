"""需求交付 service：附件文件夹管理、DDD 用户故事生成、需求分析说明书生成。

所有路径基于 Obsidian vault 派生（config.REQUIREMENT_ATTACHMENT_DIR /
REQUIREMENT_DOC_DIR），与知识库同源，便于 Obsidian 直接索引。
"""
from __future__ import annotations

import json
import os
import re
import shutil
from typing import Any, Dict, List, Optional

from sqlalchemy import func

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from db.models import PmwbUserStory, PmwbRequirementEvaluation, PmwbRequirementExt, SentEmail
from core.config import settings
from core.exceptions import NotFoundException, ValidationException
from services.storygen_rules import split_into_user_stories


# ---------------------------------------------------------------------------
# 责任部门推导（替代模板写死的「市场经营部」）
# 优先按「实际提单人归属部门」（组织分工表），其次按需求涉及系统，最后回退默认。
# 注：通讯录 SaInfo 无部门字段，提单人→部门采用组织分工表固化映射。
# ---------------------------------------------------------------------------
SYSTEM_DEPT_MAP = [
    ("专线", "政企客户部"),
    ("短信", "政企客户部"),
    ("卫星", "政企客户部"),
    ("国铁", "政企客户部"),
    ("订单中心", "订单中心"),
    ("订单", "订单中心"),
    ("电子协议", "政企客户部"),
    ("协议", "政企客户部"),
    ("CRM", "CRM"),
    ("BOSS", "BOSS"),
    ("boss", "BOSS"),
    ("生产运营", "生产运营平台"),
    ("运营平台", "生产运营平台"),
    ("经分", "数据资产领域"),
]
PROPOSER_DEPT_MAP = {
    "邵建": "政企客户部", "张舒明": "政企客户部", "张振": "政企客户部",
    "戴燕": "政企客户部", "黄何": "政企客户部", "金韡": "政企客户部",
    "顾宏明": "政企客户部", "戴晓飞": "生产运营平台", "秦新": "政企客户部",
    "郑文东": "CRM", "陈增明": "BOSS", "王辅松": "订单中心",
    "吴雨霜": "CRM", "叶振宇": "CRM", "张茜": "CRM", "李蕊": "BOSS",
    "陈山": "订单中心", "方舟": "政企客户部",
}
DEFAULT_DEPT = "政企客户部"


def _resolve_responsible_dept(item, db=None) -> str:
    """根据需求实际提出人归属部门推导责任部门（需求提出单位）。

    优先级：①人员主数据(pmwb_staff)归属组织 → ②提单人固化映射(fallback) →
    ③需求涉及系统归属部门 → ④提单人名中的系统线索 → ⑤默认。
    """
    system = (getattr(item, "system_name", "") or "") if item else ""
    proposer = (getattr(item, "proposer", "") or "") if item else ""
    # ① 人员主数据：提单人 → 所属组织（基础数据统一维护）
    if proposer and db is not None:
        try:
            from services.basic_data import basic_data_service

            dept = basic_data_service.dept_of(db, proposer)
            if dept:
                return dept
        except Exception:  # noqa: BLE001 主数据查询失败时回退固化映射
            pass
    # ② 提单人固化归属部门（fallback，主数据未覆盖时兜底）
    if proposer and proposer in PROPOSER_DEPT_MAP:
        return PROPOSER_DEPT_MAP[proposer]
    # ② 需求涉及系统 → 归属部门
    for kw, dept in SYSTEM_DEPT_MAP:
        if kw and kw in system:
            return dept
    # ③ 提单人名中的系统线索
    for kw, dept in SYSTEM_DEPT_MAP:
        if kw and kw in proposer:
            return dept
    return DEFAULT_DEPT


# 每 20 人天拆分一个用户故事（拆分参考粒度）
WORKLOAD_PER_STORY = 20.0


# ---------------------------------------------------------------------------
# 路径与文件夹
# ---------------------------------------------------------------------------
def _safe_name(name: str) -> str:
    """把需求标题转成安全的文件夹名片段。"""
    if not name:
        return "untitled"
    s = re.sub(r'[\\/:*?"<>|\r\n\t]+', "_", name).strip()
    s = re.sub(r'\s+', "_", s)
    return s[:40]


def _requirement_base(req_id: str) -> str:
    """需求级根目录名：REQ-XXXX_标题。"""
    item = None
    # 延迟导入避免循环；这里只取值
    return req_id


def _resolve_paths(req_id: str, req_name: Optional[str] = None) -> Dict[str, str]:
    """需求归档统一目录：所有附件与生成文档均放在「需求分析说明书」下同一子文件夹。"""
    vault = settings.OBSIDIAN_VAULT_PATH
    folder_base = os.path.join(vault, settings.REQUIREMENT_DOC_DIR)
    sub = f"{req_id}_{_safe_name(req_name or req_id)}"
    folder = os.path.join(folder_base, sub)
    # 旧附件目录（若存在则迁移）
    old_att_base = os.path.join(vault, settings.REQUIREMENT_ATTACHMENT_DIR)
    old_att_folder = os.path.join(old_att_base, sub)
    return {
        "folder_base": folder_base,
        "folder": folder,
        "old_att_folder": old_att_folder,
    }


def init_folder(db, req_id: str) -> Dict[str, Any]:
    """创建需求分析说明书文件夹（幂等），并迁移旧附件目录，返回路径与当前文件列表。"""
    item = db.query(SentEmail).filter(SentEmail.req_id == req_id).first()
    req_name = item.req_name if item else None
    paths = _resolve_paths(req_id, req_name)
    folder = paths["folder"]
    old_att_folder = paths["old_att_folder"]
    os.makedirs(folder, exist_ok=True)
    # 兼容迁移：旧「需求附件」文件夹存在时，把文件搬到新统一目录
    if os.path.isdir(old_att_folder) and not os.path.samefile(old_att_folder, folder):
        for fn in os.listdir(old_att_folder):
            src = os.path.join(old_att_folder, fn)
            dst = os.path.join(folder, fn)
            if os.path.isfile(src) and not os.path.exists(dst):
                shutil.move(src, dst)
    attachments = _list_attachments_from(folder)
    # 环节时间日志：打开需求工作流 → 记录「需求采集」环节进入时间（取最早邮件创建时间）
    try:
        from services.requirement_stage import record_stage_entered

        earliest = (
            db.query(SentEmail.created_at)
            .filter(SentEmail.req_id == req_id)
            .order_by(SentEmail.id.asc())
            .first()
        )
        record_stage_entered(db, req_id, "collect", at=earliest[0] if earliest else None)
    except Exception:  # noqa: BLE001
        pass
    return {
        "req_id": req_id,
        "folder": folder,
        "attachments": attachments,
    }


def _human_size(num: int) -> str:
    for unit in ("B", "KB", "MB", "GB"):
        if num < 1024:
            return f"{num:.0f} {unit}" if unit == "B" else f"{num:.1f} {unit}"
        num /= 1024
    return f"{num:.1f} TB"


def _list_attachments_from(folder: str) -> List[Dict[str, Any]]:
    if not os.path.isdir(folder):
        return []
    out = []
    for fn in sorted(os.listdir(folder)):
        fp = os.path.join(folder, fn)
        if os.path.isfile(fp):
            sz = os.path.getsize(fp)
            out.append({"name": fn, "bytes": sz, "size": _human_size(sz)})
    return out


def list_attachments(db, req_id: str) -> List[Dict[str, Any]]:
    item = db.query(SentEmail).filter(SentEmail.req_id == req_id).first()
    paths = _resolve_paths(req_id, item.req_name if item else None)
    return _list_attachments_from(paths["folder"])


def upload_attachment(db, req_id: str, filename: str, content: bytes) -> Dict[str, Any]:
    item = db.query(SentEmail).filter(SentEmail.req_id == req_id).first()
    paths = _resolve_paths(req_id, item.req_name if item else None)
    os.makedirs(paths["folder"], exist_ok=True)
    safe_fn = re.sub(r'[\\/:*?"<>|\r\n\t]+', "_", filename)
    fp = os.path.join(paths["folder"], safe_fn)
    with open(fp, "wb") as f:
        f.write(content)
    rel_local = os.path.relpath(fp, settings.OBSIDIAN_VAULT_PATH).replace("\\", "/")
    return {"name": safe_fn, "bytes": len(content), "size": _human_size(len(content)), "path": rel_local}


def upload_requirement_manual(
    db,
    req_id: str,
    filename: str,
    content: bytes,
    note: str = "操作手册",
) -> Dict[str, Any]:
    """上传操作手册并自动归档到业务知识交付物目录、同步主笔记。

    流程：
    1. 校验需求状态为 closed 且已设置 domain_code。
    2. 文件先落到需求分析说明书文件夹（作为源文件）。
    3. 登记到 ext.deliverables JSON。
    4. 调用 archive_requirement_manual 复制到 05-交付物/attachments/。
    5. 触发 sync_main_note_from_links 刷新主笔记 §6 内链。
    """
    ext = db.query(PmwbRequirementExt).filter(PmwbRequirementExt.req_id == req_id).first()
    if not ext:
        raise NotFoundException(f"需求不存在：{req_id}")
    if (ext.status or "proposed") not in ("accepted", "dev", "closed"):
        raise ValidationException("操作手册在需求「已采纳/开发中/已上线」阶段可上传")
    if not ext.domain_code:
        raise ValidationException("需求未设置业务领域，无法归档操作手册")

    # 1. 落盘到需求分析说明书文件夹
    upload_info = upload_attachment(db, req_id, filename, content)
    rel_local = upload_info["path"]

    # 2. 登记到 deliverables JSON
    from services.obsidian_link import add_requirement_deliverable, archive_requirement_manual
    entry = add_requirement_deliverable(db, req_id, upload_info["name"], rel_local, note or "操作手册")

    # 3. 归档到业务知识交付物目录（会更新 entry.archived_at / obsidian_path）
    archive_result = archive_requirement_manual(db, req_id)

    # 4. 触发主笔记 §6 同步，生成 Obsidian 内链
    from services.knowledge_link_service import sync_main_note_from_links
    sync_result = sync_main_note_from_links(db, ext.domain_code)

    # 取归档后的 obsidian_path（archive_requirement_manual 会回填 deliverables 数组）
    obsidian_path = entry.get("obsidian_path")
    if not obsidian_path and archive_result.get("archived"):
        obsidian_path = archive_result["archived"][-1].get("obsidian_path")

    return {
        "req_id": req_id,
        "file_name": upload_info["name"],
        "local_path": rel_local,
        "obsidian_path": obsidian_path or "",
        "archived": bool(archive_result.get("archived")),
        "main_note": archive_result.get("main_note"),
        "main_note_synced": sync_result.get("changed", False),
    }


def delete_attachment(db, req_id: str, filename: str) -> bool:
    item = db.query(SentEmail).filter(SentEmail.req_id == req_id).first()
    paths = _resolve_paths(req_id, item.req_name if item else None)
    fp = os.path.join(paths["folder"], os.path.basename(filename))
    # 防止路径穿越：必须落在统一文件夹内
    if not os.path.abspath(fp).startswith(os.path.abspath(paths["folder"])):
        return False
    if os.path.isfile(fp):
        os.remove(fp)
        return True
    return False


# ---------------------------------------------------------------------------
# DDD 用户故事生成
# ---------------------------------------------------------------------------
# 系统/标题关键词 -> DDD 子域映射（领域层体现）
_DDD_MAP = [
    ("专线", "政企专线领域", "国际版专线受理"),
    ("短信", "集团短信领域", "短信实名制"),
    ("宽带", "家宽/宽带领域", "流失预警与挽回"),
    ("网关", "网络接入领域", "被动销户治理"),
    ("CRM", "客户经营领域", "客户受理与工单"),
    ("BOSS", "计费账务领域", "计费与用户状态"),
    ("数据", "数据资产领域", "经分与模型"),
    ("运营", "生产运营领域", "运营监控与保障"),
]


def _derive_ddd(req_name: str, system_name: Optional[str]) -> Dict[str, str]:
    text = f"{req_name or ''} {system_name or ''}"
    for kw, domain, sub in _DDD_MAP:
        if kw in text:
            return {
                "domain": domain,
                "subdomain": sub,
                "aggregate": "需求-评估-交付",
                "entity": "需求、用户故事、开发工单",
            }
    return {
        "domain": "政企需求交付",
        "subdomain": "需求评估与履约",
        "aggregate": "需求-评估-交付",
        "entity": "需求、用户故事、开发工单",
    }


def _coarse_features(text: str) -> List[str]:
    """把澄清内容按「行」拆为粗粒度功能块（不过度切分长句）。

    每行视为一个相对独立的功能诉求；仅去掉前导编号与过短噪声。
    这样一条功能块 ≈ 一个用户故事候选，避免原逻辑按句号切碎导致粒度过细。
    """
    if not text or not text.strip():
        return []
    lines = [l.strip() for l in text.replace("\r", "\n").split("\n")]
    feats: List[str] = []
    for l in lines:
        # 去前导编号 1. 1、 (1) ①
        l = re.sub(r"^[（(]?\d+[.、)）]?\s*", "", l).strip()
        if len(l) < 4:
            continue
        feats.append(l)
    # 去重并保持顺序
    seen = set()
    uniq = []
    for f in feats:
        if f not in seen:
            seen.add(f)
            uniq.append(f)
    return uniq


def _chunk(items: List[str], n: int) -> List[List[str]]:
    """将 items 顺序均匀分成 n 组（前 m 组多 1 个），空组剔除。"""
    if not items:
        return []
    n = max(1, min(n, len(items)))
    k, m = divmod(len(items), n)
    groups, idx = [], 0
    for i in range(n):
        size = k + (1 if i < m else 0)
        if size <= 0:
            continue
        groups.append(items[idx: idx + size])
        idx += size
    return [g for g in groups if g]


def _total_workload(db, req_id: str) -> float:
    """汇总需求团队评估总工作量（人天）。"""
    rows = (
        db.query(PmwbRequirementEvaluation)
        .filter(PmwbRequirementEvaluation.req_id == req_id)
        .all()
    )
    return float(sum(float((r.workload or 0)) for r in rows))


def _story_to_dict(st: PmwbUserStory) -> Dict[str, Any]:
    return {
        "id": st.id,
        "seq": st.seq,
        "title": st.title or "",
        "desc": st.desc or "",
        "scene": st.scene or "",
        "acceptance": json.loads(st.acceptance) if st.acceptance else [],
        "rules": json.loads(st.rules) if st.rules else [],
        "finalized": bool(st.finalized),
    }


def get_user_stories(db, req_id: str) -> Dict[str, Any]:
    """从数据库读取需求下的用户故事列表。"""
    rows = (
        db.query(PmwbUserStory)
        .filter(PmwbUserStory.req_id == req_id)
        .order_by(PmwbUserStory.seq.asc())
        .all()
    )
    return {"req_id": req_id, "stories": [_story_to_dict(r) for r in rows]}


def search_user_stories(
    db,
    keyword: Optional[str] = None,
    finalized: Optional[int] = None,
    page: int = 1,
    page_size: int = 20,
) -> Dict[str, Any]:
    """全局用户故事模糊查询（跨需求）。

    - keyword：空格分词，多词 AND；每个词在 标题/描述/场景/验收标准/业务规则/
      需求编号/需求名称 全字段 OR 模糊匹配。
    - finalized：None=全部，0=草稿，1=已定稿。
    - 排序：created_at 倒序（同刻按 id 倒序）。
    - 分页：page / page_size。
    返回 {items, total, page, page_size}，item 含 req_name。
    """
    from sqlalchemy import and_, func, or_

    # 每个 req_id 取一条需求名称（团队多行时取任一），便于关联展示与检索
    sub = (
        db.query(
            SentEmail.req_id.label("req_id"),
            func.min(SentEmail.req_name).label("req_name"),
        )
        .group_by(SentEmail.req_id)
        .subquery()
    )
    q = db.query(PmwbUserStory, sub.c.req_name).outerjoin(
        sub, PmwbUserStory.req_id == sub.c.req_id
    )

    if finalized is not None:
        q = q.filter(PmwbUserStory.finalized == (1 if finalized else 0))

    words = [w for w in re.split(r"\s+", (keyword or "").strip()) if w]
    for w in words:
        like = f"%{w}%"
        q = q.filter(
            or_(
                PmwbUserStory.title.like(like),
                PmwbUserStory.desc.like(like),
                PmwbUserStory.scene.like(like),
                PmwbUserStory.acceptance.like(like),
                PmwbUserStory.rules.like(like),
                PmwbUserStory.req_id.like(like),
                sub.c.req_name.like(like),
            )
        )

    total = q.count()
    q = q.order_by(PmwbUserStory.created_at.desc(), PmwbUserStory.id.desc())
    page = max(1, page)
    page_size = max(1, min(page_size, 200))
    rows = q.offset((page - 1) * page_size).limit(page_size).all()

    items: List[Dict[str, Any]] = []
    for st, req_name in rows:
        d = _story_to_dict(st)
        d["req_id"] = st.req_id
        d["req_name"] = req_name or ""
        d["created_at"] = st.created_at.isoformat() if st.created_at else None
        items.append(d)

    return {"items": items, "total": total, "page": page, "page_size": page_size}


def get_user_story_stats(db) -> Dict[str, int]:
    """用户故事全局统计：总数、已定稿、草稿。"""
    total = db.query(func.count(PmwbUserStory.id)).scalar() or 0
    finalized = (
        db.query(func.count(PmwbUserStory.id)).filter(PmwbUserStory.finalized == 1).scalar() or 0
    )
    draft = (
        db.query(func.count(PmwbUserStory.id)).filter(PmwbUserStory.finalized == 0).scalar() or 0
    )
    return {"total": total, "finalized": finalized, "draft": draft}


def save_user_stories(db, req_id: str, stories: List[Dict[str, Any]]) -> Dict[str, Any]:
    """全量替换需求下的用户故事。"""
    # 删除旧记录
    db.query(PmwbUserStory).filter(PmwbUserStory.req_id == req_id).delete()
    # 插入新记录
    for i, s in enumerate(stories, start=1):
        st = PmwbUserStory(
            req_id=req_id,
            seq=i,
            title=s.get("title", ""),
            desc=s.get("desc", ""),
            scene=s.get("scene", ""),
            acceptance=json.dumps(s.get("acceptance") or [], ensure_ascii=False),
            rules=json.dumps(s.get("rules") or [], ensure_ascii=False),
            finalized=1 if s.get("finalized") else 0,
        )
        db.add(st)
    db.commit()
    # 环节时间日志：用户故事首次落库 → 记录「用户故事」环节进入时间
    if stories:
        try:
            from services.requirement_stage import record_stage_entered

            record_stage_entered(db, req_id, "story")
        except Exception:  # noqa: BLE001
            pass
    return get_user_stories(db, req_id)


def generate_user_stories(
    db, req_id: str, content: str, strategy: str = "rules_v2"
) -> Dict[str, Any]:
    """基于澄清内容生成用户故事，并落库。

    strategy 支持：
    - "rules_v2"（默认）：合并优先的规则引擎 v2，符合公司最新管理规范
    - "rules_v1"：旧版按行/人天拆分（兼容模式，不推荐）
    - "llm"：AI 智能生成（复用 AI 中心统一大模型注册表，见 services.llm_provider）

    旧版兼容：不传 strategy 走 v2。
    """
    item = db.query(SentEmail).filter(SentEmail.req_id == req_id).first()
    req_name = item.req_name if item else req_id
    system_name = item.system_name if item else None
    proposer = item.proposer if item else ""
    ddd = _derive_ddd(req_name, system_name)

    # 内容优先用澄清内容，回退到原始描述
    source = (content or "").strip() or (item.description or "" if item else "")
    if not source.strip():
        source = (
            f"{req_name}：需建设对应能力，支撑政企业务运营，"
            "在相关系统中实现受理、配置、查询与监控等闭环功能。"
        )

    if strategy == "rules_v1":
        stories = _generate_v1(source, ddd, db, req_id)
        strategy_used = "rules_v1"
    elif strategy == "llm":
        # LLM 策略：尝试调用 LLM，失败回退 rules_v2
        stories, strategy_used = _generate_with_llm_fallback(source, ddd, db, req_id)
    else:
        # 默认 rules_v2
        stories = _generate_v2(source, ddd)
        strategy_used = "rules_v2"

    # 落库并返回持久化后的完整数据
    saved = save_user_stories(db, req_id, stories)
    return {
        "req_id": req_id,
        "ddd": ddd,
        "proposer": proposer,
        "stories": saved["stories"],
        "strategy_used": strategy_used,
    }


def _generate_v2(source: str, ddd: Dict[str, str]) -> List[Dict[str, Any]]:
    """规则引擎 v2：合并优先策略。"""
    return split_into_user_stories(source, ddd)


def _generate_v1(source: str, ddd: Dict[str, str], db, req_id: str) -> List[Dict[str, Any]]:
    """规则引擎 v1（旧版兼容）：按行/人天拆分。"""
    features = _coarse_features(source)
    if not features:
        features = [source]

    total_wl = _total_workload(db, req_id)
    if total_wl and total_wl > 0:
        target = max(1, round(total_wl / WORKLOAD_PER_STORY))
    else:
        target = max(1, (len(features) + 1) // 2)
    target = min(target, 8, len(features))

    groups = _chunk(features, target)
    if not groups:
        groups = [[source]]

    role = "业务负责人员"
    purpose = "支撑政企业务数字化运营，保障业务连续性与客户体验"
    stories = []
    for i, group in enumerate(groups, start=1):
        function = group[0] if len(group) == 1 else f"{group[0]}等{len(group)}项能力"
        feat_text = "；".join(group)
        title = f"US{i}：作为{role}，希望{function}，以便{purpose}"
        desc = (
            f"【故事描述】{feat_text}。该需求来源于政企业务实际运营场景，"
            f"需在「{ddd['subdomain']}」相关系统中实现上述能力，"
            "以保障业务连续性与客户体验，并支持端到端的功能验证。"
        )
        scene = (
            f"【故事场景】当业务人员处理「{group[0]}」相关操作时，系统应提供对应的功能入口与数据支撑，"
            "使其能够在一次完整流程内完成目标操作，并获得明确的结果反馈；"
            "若操作失败，应给出可定位的提示。"
        )
        acceptance = [f"验证{x}功能是否成功实现" for x in group]
        stories.append({
            "seq": i,
            "title": title,
            "desc": desc,
            "scene": scene,
            "acceptance": acceptance,
            "rules": [],
            "finalized": False,
        })
    return stories


def _generate_with_llm_fallback(
    source: str, ddd: Dict[str, str], db, req_id: str
) -> tuple:
    """AI 智能生成（复用 AI 中心统一大模型注册表），失败回退 rules_v2。

    不再单独依赖 US_STORY_LLM_* 配置；可用性以前端 llm-status 探测的
    AI 中心统一状态为准（有任一 enabled 提供方即可用）。
    """
    try:
        from services.storygen_llm import generate_via_unified
        stories = generate_via_unified(db, source, ddd)
        return stories, "llm"
    except Exception:
        # 降级到 v2
        stories = _generate_v2(source, ddd)
        return stories, "rules_v2_fallback"


# ---------------------------------------------------------------------------
# 需求分析说明书生成
# ---------------------------------------------------------------------------
def _find_heading_index(paras, text, style_contains: str) -> int:
    for i, p in enumerate(paras):
        if text in p.text and style_contains in p.style.name:
            return i
    return -1


def _insert_after(anchor: Paragraph, text: str, style_name: Optional[str] = None, doc=None) -> Paragraph:
    """在 anchor 段落之后插入一个新段落并返回其对象（用于链式插入）。

    style_name 存在时应用段落样式；doc 必须传入 Document 对象以便查 styles
    （anchor._parent 是 _Body，无 .styles 属性，不能直接用）。
    """
    new_p = OxmlElement("w:p")
    anchor._element.addnext(new_p)
    para = Paragraph(new_p, anchor._parent)
    if style_name and doc is not None:
        try:
            para.style = doc.styles[style_name]
        except KeyError:
            pass
    if text:
        para.add_run(text)
    return para


def _set_font_yh(run) -> None:
    """把 run 字体设为微软雅黑（含中文 eastAsia / ascii / hAnsi）。"""
    run.font.name = "微软雅黑"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        rfonts.set(qn(attr), "微软雅黑")


def _style_runs_yh(para) -> None:
    """对段落所有 run 统一设置微软雅黑字体。"""
    if para is None:
        return
    for r in para.runs:
        _set_font_yh(r)


def _insert_table_after(anchor, rows: int, cols: int, doc):
    """在 anchor 段落之后插入一张表格并返回其对象。"""
    tbl = doc.add_table(rows=rows, cols=cols)
    anchor._element.addnext(tbl._element)
    return tbl


def _tables_between(doc, start_el, end_el):
    """返回文档正文中位于 start_el 与 end_el 之间的所有表格元素（按文档顺序）。"""
    body = doc.element.body
    res, collecting = [], False
    for ch in body.iterchildren():
        if ch is start_el:
            collecting = True
            continue
        if ch is end_el:
            break
        if collecting and ch.tag == qn("w:tbl"):
            res.append(ch)
    return res


# 规则表表头（与模板《需求分析说明书》一致：7 列，含规则编号双列）
_RULE_HEADERS = ["用户故事", "规则编号", "规则编号", "领域", "规则描述", "举例", "影响的主要功能"]


def _insert_rules_table(anchor, doc, story_title: str, rules):
    """在每个用户故事下插入规则表（模板要求：表可空但必须有）。"""
    rules = rules or []
    nrows = max(1, len(rules))  # 至少 1 行占位，保证表格存在
    tbl = _insert_table_after(anchor, nrows + 1, len(_RULE_HEADERS), doc)
    # 表头
    for ci, h in enumerate(_RULE_HEADERS):
        cell = tbl.rows[0].cells[ci]
        cell.text = h
        _style_runs_yh(cell.paragraphs[0])
    # 规则行（无规则则留一个占位行）
    for ri in range(1, nrows + 1):
        row = tbl.rows[ri]
        rtext = rules[ri - 1] if ri - 1 < len(rules) else ""
        rid = f"R{ri:02d}"
        vals = [story_title, rid, rid, "—", rtext, "—", "—"]
        for ci, v in enumerate(vals):
            cell = row.cells[ci]
            cell.text = v
            _style_runs_yh(cell.paragraphs[0])
    return tbl


def generate_doc(db, req_id: str, stories: List[Dict[str, Any]], clarification: str = "") -> Dict[str, Any]:
    """基于固定模板生成《需求分析说明书》，仅填充第1/2/3章，其余复用模板。

    第1章 基本信息（表格）：需求单编号、需求责任部门（按实际归属推导）、需求提出人
    第2章 用户故事-背景以及业务总体描述：原始需求内容/澄清
    第3章 用户故事-用户Story：自动填充 US1/US2...（固定4段模板），每个故事下补规则表
    补充信息（含表格单元格）统一使用微软雅黑字体
    第4章 需求检查项 / 第5章 版本历史：保持模板原样
    """
    item = db.query(SentEmail).filter(SentEmail.req_id == req_id).first()
    req_name = item.req_name if item else req_id
    proposer = item.proposer if item else ""
    paths = _resolve_paths(req_id, req_name)
    os.makedirs(paths["folder"], exist_ok=True)

    template = settings.REQUIREMENT_DOC_TEMPLATE
    if not os.path.isfile(template):
        raise FileNotFoundError(f"模板不存在：{template}")
    doc = Document(template)
    paras = doc.paragraphs

    # ---- 第1章 基本信息（templates 首张表）----
    if doc.tables:
        tbl = doc.tables[0]
        if len(tbl.rows) >= 2:
            c0 = tbl.rows[1].cells[0]   # 需求单编号
            c1 = tbl.rows[1].cells[1]   # 需求提出单位 / 责任部门
            c2 = tbl.rows[1].cells[2]   # 需求提出人
            c0.text = req_id
            c1.text = _resolve_responsible_dept(item, db=db)
            c2.text = proposer or ""
            for c in (c0, c1, c2):
                _style_runs_yh(c.paragraphs[0])

    # ---- 第2章 背景以及业务总体描述 ----
    idx_bg = _find_heading_index(paras, "背景以及业务总体描述", "Heading 3")
    if idx_bg >= 0 and idx_bg + 1 < len(paras):
        bg_body = paras[idx_bg + 1]
        bg_text = (clarification or "").strip() or (item.description or "" if item else "")
        if not bg_text:
            bg_text = req_name
        bg_body.text = f"【原始需求内容】{bg_text}"
        _style_runs_yh(bg_body)

    # ---- 第3章 用户Story：清除旧 US 区块（含模板残留规则表），按生成结果填充 ----
    idx_story = _find_heading_index(paras, "用户Story", "Heading 3")
    idx_check = _find_heading_index(paras, "需求检查项", "Heading 2")
    if idx_story >= 0 and idx_check > idx_story:
        # 删除 idx_story+1 .. idx_check-1 之间的旧段落（原模板 US1/US2 及其正文）
        old = [paras[i] for i in range(idx_story + 1, idx_check)]
        for p in old:
            p._element.getparent().remove(p._element)
        # 移除该区间内模板残留的规则表（table 元素不被段落删除逻辑清除）
        for t in _tables_between(doc, paras[idx_story]._element, paras[idx_check]._element):
            t.getparent().remove(t)
        # 链式插入新故事 + 每个故事下的规则表
        anchor = paras[idx_story]
        for st in stories:
            anchor = _insert_after(anchor, st.get("title", ""), "Heading 4", doc)
            _style_runs_yh(anchor)
            anchor = _insert_after(anchor, f"Story功能描述\n{st.get('desc', '')}", None, doc)
            _style_runs_yh(anchor)
            anchor = _insert_after(anchor, f"Story场景\n{st.get('scene', '')}", None, doc)
            _style_runs_yh(anchor)
            anchor = _insert_after(anchor, "验收标准\n" + "；".join(st.get("acceptance") or []), None, doc)
            _style_runs_yh(anchor)
            # 每个故事下插入规则表（模板要求：表可空但必须有）
            anchor = _insert_rules_table(anchor, doc, st.get("title", ""), st.get("rules"))

    # ---- 落盘 ----
    out_name = f"{req_name or req_id}需求分析说明书.docx"
    out_path = os.path.join(paths["folder"], out_name)
    doc.save(out_path)

    rel = os.path.relpath(out_path, settings.OBSIDIAN_VAULT_PATH)
    # 环节时间日志：首次成功生成需求分析说明书 → 记录「生成文档」环节进入时间
    try:
        from services.requirement_stage import record_stage_entered

        record_stage_entered(db, req_id, "doc")
    except Exception:  # noqa: BLE001
        pass
    url = "obsidian://open?path=" + out_path
    return {
        "req_id": req_id,
        "file": out_name,
        "path": out_path,
        "url": url,
        "relative_path": rel,
    }
